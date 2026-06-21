# -*- coding: utf-8 -*-
"""国民健康保険税（料）減免申請書 サンプル書類ジェネレータ。

PDF5種（適正/要改善/重大/災害/死亡）＋作成用Wordひな形を生成する。
内容は各自治体の減免取扱要綱の一般的な基準に基づく架空のサンプル。

    python generate_genmen_docs.py
"""
from __future__ import annotations

from pathlib import Path

from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from docx import Document
from docx.shared import Pt

HERE = Path(__file__).resolve().parent

FONT_PATHS = [
    Path(r"C:\Windows\Fonts\meiryo.ttc"),
    Path(r"C:\Windows\Fonts\msgothic.ttc"),
    Path(r"C:\Windows\Fonts\YuGothM.ttc"),
]


def register_font() -> str:
    for fp in FONT_PATHS:
        if fp.exists():
            name = fp.stem.replace(" ", "")
            pdfmetrics.registerFont(TTFont(name, str(fp)))
            return name
    raise FileNotFoundError("日本語フォントが見つかりません")


def esc(text: str) -> str:
    return (
        str(text)
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace("\n", "<br/>")
    )


def build_pdf(out: Path, title: str, intro: str, rows: list[tuple[str, str]]):
    font = register_font()
    ss = getSampleStyleSheet()
    title_style = ParagraphStyle("t", parent=ss["Title"], fontName=font, fontSize=15, leading=20, alignment=TA_LEFT, spaceAfter=4 * mm)
    intro_style = ParagraphStyle("i", parent=ss["BodyText"], fontName=font, fontSize=9, leading=13, spaceAfter=4 * mm, textColor=colors.HexColor("#444444"))
    label_style = ParagraphStyle("l", parent=ss["BodyText"], fontName=font, fontSize=9, leading=12)
    body_style = ParagraphStyle("b", parent=ss["BodyText"], fontName=font, fontSize=9, leading=13)

    story: list = [Paragraph(esc(title), title_style)]
    if intro:
        story.append(Paragraph(esc(intro), intro_style))

    data = [[Paragraph(esc(lbl), label_style), Paragraph(esc(val), body_style)] for lbl, val in rows]
    table = Table(data, colWidths=[48 * mm, 132 * mm])
    table.setStyle(
        TableStyle(
            [
                ("BOX", (0, 0), (-1, -1), 0.7, colors.black),
                ("INNERGRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#999999")),
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#eef2f7")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    story.append(table)
    story.append(Spacer(1, 2 * mm))

    doc = SimpleDocTemplate(
        str(out), pagesize=A4,
        leftMargin=15 * mm, rightMargin=15 * mm, topMargin=14 * mm, bottomMargin=14 * mm,
        title=title, author="源内AI サンプル",
    )
    doc.build(story)
    print("created:", out.name)


# ============================================================
# サンプル定義
# ============================================================
COMMON_NOTE = "※これは源内AI「国保税減免申請書 審査」アプリの動作確認用サンプルです（架空の事例）。実在の個人・団体とは関係ありません。"

# ① 非自発的失業（会社都合解雇）— 適正（A評価想定）
S1 = (
    HERE / "genmen_1_shitsugyo.pdf",
    "国民健康保険税減免（軽減）申請書",
    COMMON_NOTE + "【区分：非自発的失業者の軽減 / 想定：適正】",
    [
        ("申請日 / あて先", "令和8年5月20日 / 緑川市長 あて"),
        ("世帯主氏名 / 住所", "田中 一郎（昭和55年4月3日生・44歳） / 緑川市中央1-2-3"),
        ("被保険者記号番号 / 連絡先", "記号 1234 番号 56789 / 電話 090-0000-0001"),
        ("対象被保険者", "田中 一郎（世帯主・本人）、田中 花子（妻・39歳・無職）"),
        ("申請区分（減免事由）", "非自発的失業者の軽減（雇用保険の特定受給資格者）"),
        ("事由の具体的内容", "勤務先（株式会社みどり製作所）の事業縮小に伴う人員整理により、令和8年4月30日付で会社都合により解雇された。再就職活動中で現在無収入。"),
        ("離職理由 / 離職日", "雇用保険受給資格者証の離職理由コード「11（解雇）」 / 令和8年4月30日離職、5月1日に国保加入（資格取得から14日以内）"),
        ("前年（令和7年）給与所得", "給与収入 420万円（給与所得 296万円）"),
        ("軽減の内容（申請）", "前年給与所得を100分の30（約88.8万円相当）とみなして算定する軽減を希望。軽減期間：離職日の翌日（令和8年5月1日）〜令和9年度末。"),
        ("添付書類", "①雇用保険受給資格者証（写し）②資格確認書 ③本人確認書類（運転免許証写し） ※すべて添付済み"),
        ("他制度の状況", "法定軽減（7割5割2割）は前年所得超過のため非該当。他保険への加入予定なし。"),
        ("申請者署名・押印", "田中 一郎 ㊞"),
    ],
)

# ② 所得激減（事業不振）— 要改善（添付・記載不足）
S2 = (
    HERE / "genmen_2_shotokugekigen.pdf",
    "国民健康保険税減免申請書",
    COMMON_NOTE + "【区分：所得激減（事業不振） / 想定：要改善】",
    [
        ("申請日 / あて先", "令和8年6月18日 / 緑川市長 あて"),
        ("世帯主氏名 / 住所", "鈴木 健二（51歳） / 緑川市東町5-6"),
        ("被保険者記号番号", "記号 2233 番号 44556"),
        ("対象被保険者", "鈴木 健二（世帯主・自営業）、鈴木 美和（妻・パート）"),
        ("申請区分（減免事由）", "所得激減（事業の不振）"),
        ("事由の具体的内容", "営む飲食店の売上が取引先の撤退により急減し、令和8年に入り収入が大幅に減少した。"),
        ("所得の減少状況", "前年（令和7年）事業所得 約240万円。本年見込みは月平均で前年比およそ3〜4割程度に落ち込む見通し。"),
        ("減免希望内容", "所得割額の一部減免を希望。"),
        ("添付書類", "①確定申告書（令和7年分）の写し のみ添付。※給与・売上の減少を示す月次資料、預貯金残高、収入状況報告書は未添付。"),
        ("他制度の状況", "記載なし（法定軽減・非自発的失業軽減との関係について記載がない）。"),
        ("申請者署名・押印", "鈴木 健二（押印なし）"),
    ],
)

# ③ 対象外事由・期限超過 — 重大（却下リスク）
S3 = (
    HERE / "genmen_3_taisyogai.pdf",
    "国民健康保険税減免申請書",
    COMMON_NOTE + "【区分：所得激減として申請されているが対象外の疑い / 想定：重大な不備】",
    [
        ("申請日 / あて先", "令和8年9月25日 / 緑川市長 あて"),
        ("世帯主氏名 / 住所", "高橋 三郎（60歳） / 緑川市西1-1"),
        ("被保険者記号番号", "記号 7788 番号 99001"),
        ("対象被保険者", "高橋 三郎（世帯主）"),
        ("申請区分（減免事由）", "所得激減として申請"),
        ("事由の具体的内容", "本年の所得が前年より減った。前年は保有していた上場株式の売却益（譲渡所得）約500万円と株式配当があり所得が大きかったが、本年はその譲渡・配当がないため所得が減少した。"),
        ("所得の減少状況", "前年（令和7年）合計所得 約560万円（うち株式譲渡所得500万円・配当所得40万円）。本年見込み 給与・年金等で約180万円。"),
        ("就労・収入の状況", "本年7月より再就職が決定し、8月から会社の健康保険に加入予定（国保は喪失予定）。現在の収入は安定して回復している。"),
        ("申請対象の保険税", "第1期〜第3期分（いずれも納期限を経過。第1期は納付済み）。"),
        ("減免希望内容", "保険税の全額減免を希望。"),
        ("添付書類", "令和7年分確定申告書の写し。"),
        ("申請者署名・押印", "高橋 三郎 ㊞"),
    ],
)

# ④ 災害（火災）— 適正（A評価想定）
S4 = (
    HERE / "genmen_4_saigai.pdf",
    "国民健康保険税減免申請書",
    COMMON_NOTE + "【区分：災害（火災） / 想定：適正】",
    [
        ("申請日 / あて先", "令和8年8月5日 / 緑川市長 あて"),
        ("世帯主氏名 / 住所", "佐々木 良子（48歳） / 緑川市南町8-9"),
        ("被保険者記号番号 / 連絡先", "記号 5566 番号 11223 / 電話 090-0000-0004"),
        ("対象被保険者", "佐々木 良子（世帯主）、佐々木 太一（長男・16歳）"),
        ("申請区分（減免事由）", "災害（住宅の火災による損害）"),
        ("事由の具体的内容", "令和8年7月20日に発生した自宅の火災により、住宅および家財が焼損した。"),
        ("損害の程度", "消防署の罹災証明書において、住宅・家財に対する損害割合は「全焼（10分の5以上）」と認定。"),
        ("世帯の所得", "前年（令和7年）合計所得 約310万円（500万円以下）。"),
        ("損害保険による補填", "火災保険の支払いは家財の一部のみで、住宅損害の補填には満たない（保険金支払通知の写しを添付）。"),
        ("減免希望内容", "損害割合10分の5以上・所得500万円以下に該当するため、当該年度の保険税の全部減免を希望。"),
        ("添付書類", "①罹災証明書（損害割合記載）②損害保険金支払通知の写し ③資格確認書 ④納入通知書 ※すべて添付済み"),
        ("申請者署名・押印", "佐々木 良子 ㊞"),
    ],
)

# ⑤ 死亡・疾病 — 要改善（収入・資産状況の記載不足）
S5 = (
    HERE / "genmen_5_shibou.pdf",
    "国民健康保険税減免申請書",
    COMMON_NOTE + "【区分：死亡・重篤な疾病等 / 想定：要改善】",
    [
        ("申請日 / あて先", "令和8年7月10日 / 緑川市長 あて"),
        ("世帯主氏名 / 住所", "中村 京子（58歳） / 緑川市北4-5"),
        ("被保険者記号番号", "記号 3344 番号 22110"),
        ("対象被保険者", "中村 京子（世帯主）、中村 さくら（二女・大学生・21歳）"),
        ("申請区分（減免事由）", "世帯の生計中心者の死亡"),
        ("事由の具体的内容", "世帯の生計を主に支えていた夫（中村 武・享年61）が令和8年6月15日に病気により死亡し、世帯の収入が大きく減少した。"),
        ("死亡を証する書類", "死亡診断書の写しを添付。"),
        ("世帯の収入・資産状況", "（記載が「収入が減って苦しい」のみで、現在の世帯収入額・遺族年金等の見込み・預貯金等の資産状況の具体的記載がない）"),
        ("減免希望内容", "保険税の減免を希望（割合の記載なし）。"),
        ("添付書類", "①死亡診断書の写し のみ。収入状況報告書・預貯金残高等は未添付。"),
        ("申請者署名・押印", "中村 京子 ㊞"),
    ],
)


def build_word_template(out: Path):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Meiryo"
    style.font.size = Pt(10.5)

    h = doc.add_heading("国民健康保険税（料）減免申請書（作成用ひな形）", level=0)
    p = doc.add_paragraph("※下表の空欄に記入のうえ、源内AI「国保税減免申請書 審査」アプリでテキストを貼り付けて審査できます。減免事由に応じた添付書類を必ずご用意ください。")
    p.runs[0].font.size = Pt(9)

    fields = [
        "申請日 / あて先",
        "世帯主氏名 / 生年月日 / 住所",
        "被保険者記号番号 / 連絡先",
        "対象被保険者（氏名・続柄）",
        "申請区分（減免事由）　□災害　□所得激減　□非自発的失業　□死亡・疾病・障害　□その他",
        "事由の具体的内容",
        "所得・損害等の状況（前年比の減少率／損害割合等）",
        "減免希望内容（全部／一部・対象＝所得割／均等割等・希望割合）",
        "添付書類（離職票・受給資格者証・罹災証明・源泉徴収票・給与明細・診断書・預貯金残高 等）",
        "法定軽減・他制度との関係",
        "申請者署名・押印",
    ]
    table = doc.add_table(rows=len(fields), cols=2)
    table.style = "Table Grid"
    for i, label in enumerate(fields):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = ""
        table.rows[i].cells[0].width = Pt(170)
    doc.add_paragraph("")
    note = doc.add_paragraph("【ご注意】減免は申請に基づく審査制です。譲渡所得・株式等配当所得・一時所得の減少、再就職・他保険加入の決定、収入回復済み等は対象外となる場合があります。申請は保険税（料）の納期限までに行ってください。")
    note.runs[0].font.size = Pt(9)
    doc.save(str(out))
    print("created:", out.name)


def build_txt(out: Path, title: str, intro: str, rows: list[tuple[str, str]]):
    lines = [title, "", intro, ""]
    for lbl, val in rows:
        lines.append(f"{lbl}：{val}")
    out.write_text("\n".join(lines) + "\n", encoding="utf-8")
    print("created:", out.name)


def main():
    for idx, (out, title, intro, rows) in enumerate((S1, S2, S3, S4, S5), start=1):
        build_pdf(out, title, intro, rows)
        build_txt(HERE / f"genmen{idx}.txt", title, intro, rows)
    build_word_template(HERE / "genmen_template.docx")


if __name__ == "__main__":
    main()
