# -*- coding: utf-8 -*-
"""軽自動車税・原付手続 審査 サンプル書類ジェネレータ。

PDF5種（減免適正／減免要改善／名義変更／廃車／税止め）＋作成用Wordひな形を生成。
すべて架空の事例。PDFは説明文を入れず、タイトルと内容のみ。

    python generate_keiji_docs.py
"""
from __future__ import annotations

from pathlib import Path

from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from docx import Document
from docx.shared import Pt

HERE = Path(__file__).resolve().parent

FONT_PATHS = [
    Path(r"C:\Windows\Fonts\meiryo.ttc"),
    Path(r"C:\Windows\Fonts\msgothic.ttc"),
    Path(r"C:\Windows\Fonts\YuGothM.ttc"),
]


def register_font() -> str:
    for fp in FONT_PATHS:
        if fp.exists():
            name = fp.stem.replace(" ", "")
            pdfmetrics.registerFont(TTFont(name, str(fp)))
            return name
    raise FileNotFoundError("日本語フォントが見つかりません")


def esc(text: str) -> str:
    return (
        str(text)
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace("\n", "<br/>")
    )


def build_pdf(out: Path, title: str, rows: list[tuple[str, str]]):
    font = register_font()
    ss = getSampleStyleSheet()
    title_style = ParagraphStyle("t", parent=ss["Title"], fontName=font, fontSize=15, leading=20, alignment=TA_LEFT, spaceAfter=4 * mm)
    label_style = ParagraphStyle("l", parent=ss["BodyText"], fontName=font, fontSize=9, leading=12)
    body_style = ParagraphStyle("b", parent=ss["BodyText"], fontName=font, fontSize=9, leading=13)

    story: list = [Paragraph(esc(title), title_style)]
    data = [[Paragraph(esc(lbl), label_style), Paragraph(esc(val), body_style)] for lbl, val in rows]
    table = Table(data, colWidths=[48 * mm, 132 * mm])
    table.setStyle(
        TableStyle(
            [
                ("BOX", (0, 0), (-1, -1), 0.7, colors.black),
                ("INNERGRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#999999")),
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#eef2f7")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    story.append(table)
    story.append(Spacer(1, 2 * mm))

    doc = SimpleDocTemplate(
        str(out), pagesize=A4,
        leftMargin=15 * mm, rightMargin=15 * mm, topMargin=14 * mm, bottomMargin=14 * mm,
        title=title, author="源内AI サンプル",
    )
    doc.build(story)
    print("created:", out.name)


SAMPLES = [
    (
        HERE / "keiji_1_genmen_ok.pdf",
        "軽自動車税（種別割）減免申請書",
        [
            ("申請日 / あて先", "令和8年5月12日 / 緑川市長 あて（納期限：令和8年6月1日）"),
            ("申請者（納税義務者）", "A田 B夫（仮名・52歳） 緑川市旭町1丁目 電話090-XXXX-XXXX"),
            ("車両情報", "車両番号：緑川 580 あ 12-34 / 車台番号：ABCD1234567890123（車検証と一致）"),
            ("減免類型", "障害者等使用による減免"),
            ("対象者・手帳", "申請者本人。身体障害者手帳1級（下肢機能障害）。手帳は氏名・等級・障害名・再認定欄を含む全ページの写しを添付。"),
            ("運転者と対象者の関係", "本人運転。運転免許証（写し添付）の氏名・住所は申請者と一致。"),
            ("使用目的", "週2回の通院（緑川市民病院）と日常の買い物に使用。"),
            ("他の減免の状況", "対象者名義の普通自動車なし。自動車税種別割の減免申請なし（本車両1台のみ）。"),
            ("提出書類", "①減免申請書 ②身体障害者手帳（全ページ写し）③車検証（写し）④運転免許証（写し） ※すべて添付済み"),
            ("申請者署名・日付", "A田 B夫 ／ 令和8年5月12日（納期限内の申請）"),
        ],
    ),
    (
        HERE / "keiji_2_genmen_ng.pdf",
        "軽自動車税（種別割）減免申請書",
        [
            ("申請日 / あて先", "令和8年5月28日 / 緑川市長 あて（納期限：令和8年6月1日）"),
            ("申請者（納税義務者）", "C野 D子（仮名・47歳） 緑川市栄町3丁目"),
            ("車両情報", "車両番号：緑川 580 い 56-78 / 車台番号：EFGH9876543210987"),
            ("減免類型", "障害者等使用による減免"),
            ("対象者・手帳", "対象者は同居の長男（19歳）。療育手帳の写しは氏名・判定区分のページのみ添付（交付日・再判定欄のページなし）。判定区分の記載は「B1」。"),
            ("運転者と対象者の関係", "生計同一者（母である申請者）が運転。ただし生計同一関係を示す書類（住民票等）は未添付。"),
            ("使用目的", "長男の通所施設（週5日）への送迎。"),
            ("他の減免の状況", "記載なし（世帯内の他車両・自動車税減免の有無について記載がない）。"),
            ("提出書類", "①減免申請書 ②療育手帳（一部ページ写し）③運転免許証（写し） ※車検証（または標識交付証明書）の写しが未添付。"),
            ("申請者署名・日付", "C野 D子 ／ 日付欄が空欄。"),
        ],
    ),
    (
        HERE / "keiji_3_meigi.pdf",
        "原動機付自転車 名義変更申告書・譲渡証明書",
        [
            ("届出日 / あて先", "令和8年6月20日 / 緑川市長 あて"),
            ("旧所有者", "E藤 F男（仮名） 緑川市西町2丁目"),
            ("新所有者", "G山 H美（仮名） 緑川市東町5丁目"),
            ("車両区分", "原動機付自転車（50cc以下）"),
            ("車両情報", "標識番号：緑川市 あ 1234 / 申告書の車台番号：ZXCV0987654321"),
            ("譲渡証明", "譲渡人E藤 F男の署名あり。譲渡日：令和8年6月25日（届出日6月20日より後の日付になっている）。"),
            ("標識交付証明書", "添付あり。記載の車台番号：ZXCV0987654321（申告書と一致）。"),
            ("届出者", "新所有者の知人 I川 J太（仮名）が代理で来庁。委任状の添付なし。"),
            ("本人確認書類", "代理人I川の運転免許証は提示あり。新所有者G山の本人確認書類の写しなし。"),
            ("提出書類", "①名義変更申告書 ②譲渡証明書 ③標識交付証明書 ※委任状・新所有者の本人確認書類が不足。"),
        ],
    ),
    (
        HERE / "keiji_4_haisha.pdf",
        "原動機付自転車 廃車申告書兼標識返納書",
        [
            ("届出日 / あて先", "令和8年7月2日 / 緑川市長 あて"),
            ("所有者", "K原 L介（仮名・28歳） 緑川市浜町4丁目"),
            ("車両区分", "原動機付自転車（51cc〜90cc）"),
            ("車両情報", "標識番号：緑川市 い 5678 / 申告書の車台番号：QWER1122334455"),
            ("廃車理由", "「紛失」に丸。ただし詳細欄には「先月から見当たらない。盗まれたのかもしれない」と記載（紛失か盗難か判然としない）。"),
            ("ナンバープレート", "紛失のため返納なし。弁償届（標識弁償金の手続）は未提出。盗難の場合に必要な警察への届出（届出年月日・受理番号）の記載もなし。"),
            ("標識交付証明書", "添付なし（「見当たらない」とのこと）。"),
            ("車台番号の確認", "市保有の登録情報では当該標識番号の車台番号は QWER1122334466 であり、申告書の記載（…4455）と下2桁が一致しない。"),
            ("届出者・本人確認", "本人来庁。運転免許証提示あり。"),
            ("提出書類", "①廃車申告書兼標識返納書 のみ ※プレート・登録票・弁償届（または盗難届の受理番号）が不足。"),
        ],
    ),
    (
        HERE / "keiji_5_taxdome.pdf",
        "軽自動車税 税止め申告（報告）書",
        [
            ("申告日 / あて先", "令和8年6月30日 / 緑川市長 あて"),
            ("照会者（旧納税義務者）", "M田 N江（仮名・61歳） 緑川市中町6丁目"),
            ("車両情報", "車両番号：緑川 580 う 90-12 / 車台番号：TYUI5566778899000"),
            ("実施済み手続", "名義変更（県外の親族へ譲渡）。手続日：令和8年6月18日。"),
            ("手続先機関", "軽自動車検査協会 神奈川事務所"),
            ("証明書類", "①新旧所有者の記載がある自動車検査証（写し）②検査記録事項等証明書（写し） ※車台番号・手続日とも申告書の記載と一致。"),
            ("軽自動車税申告（報告）書", "提出あり。新所有者の住所地（県外市町村）での課税となる旨を記載。"),
            ("処理経路", "本人による税止め申告（全軽自協経由の依頼はしていないことを窓口で確認済み・二重処理なし）。"),
            ("賦課期日との関係", "手続日が令和8年6月18日のため、令和8年度分は現名義（4月1日現在の所有者＝申告者）に課税、令和9年度分から新所有者の住所地課税となる整理。"),
            ("申告者署名・日付", "M田 N江 ／ 令和8年6月30日"),
        ],
    ),
]


def build_word_template(out: Path):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Meiryo"
    style.font.size = Pt(10.5)

    doc.add_heading("軽自動車税・原付手続 受付書類（作成用ひな形）", level=0)
    p = doc.add_paragraph("※下表の空欄に記入のうえ、源内AI「軽自動車税・原付手続 審査」で業務の種類（減免申請／名義変更／廃車・標識返納／税止め）を選んで審査できます。氏名は仮名で記入してください。")
    p.runs[0].font.size = Pt(9)

    fields = [
        "申請日・届出日 / あて先",
        "申請者・届出者（氏名＝仮名・住所・連絡先）",
        "車両情報（車両番号／標識番号・車台番号・車両区分）",
        "【減免】減免類型　□障害者等使用　□福祉仕様車両　□社会福祉事業用　□その他",
        "【減免】対象者・手帳（種別・等級・全ページ写しの有無）／運転者との関係（本人・生計同一・常時介護）",
        "【名義変更】旧所有者・新所有者・譲渡日・譲渡証明の有無",
        "【廃車】廃車理由（譲渡・廃棄・盗難・紛失・転出）／プレート返納の可否・弁償届・警察届出の受理番号",
        "【税止め】実施済み手続（名義変更・廃車返納等）・手続日・手続先機関・証明書類",
        "提出書類（申請書・手帳写し・車検証／標識交付証明書・免許証・委任状・譲渡証明書等）",
        "届出者の区分（本人・代理人＝委任状・相続人）と本人確認書類",
        "署名・日付",
    ]
    table = doc.add_table(rows=len(fields), cols=2)
    table.style = "Table Grid"
    for i, label in enumerate(fields):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = ""
        table.rows[i].cells[0].width = Pt(170)
    doc.add_paragraph("")
    note = doc.add_paragraph("【ご注意】軽自動車税（種別割）は4月1日現在の所有者に課税され、月割課税・還付はありません。減免申請は納期限までに行ってください。税額の算定・最終判断は市町村の基幹システム・担当職員が行います。")
    note.runs[0].font.size = Pt(9)
    doc.save(str(out))
    print("created:", out.name)


def main():
    for out, title, rows in SAMPLES:
        build_pdf(out, title, rows)
    build_word_template(HERE / "keiji_template.docx")


if __name__ == "__main__":
    main()
